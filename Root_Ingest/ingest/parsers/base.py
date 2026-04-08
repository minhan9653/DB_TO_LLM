# This module defines the common parser contract for ingest.

# Every parser implementation converts one input file into ParsedDocument.

# Shared fallback helpers keep downstream chunk/embedding code unchanged.

# Parser-specific modules only focus on extracting raw text from files.

from __future__ import annotations
from abc import ABC, abstractmethod
from pathlib import Path
from typing import Any
from Root_Ingest.ingest.models import DocumentItem, ParsedDocument

TEXT_FILE_EXTENSIONS = {".sql", ".txt", ".md"}


class BaseParser(ABC):
    """Common interface that every parser implementation must follow."""

    parser_name = "base"

    def __init__(self, options: dict[str, Any] | None = None) -> None:
        """
        역할:
        INGEST 기본 파서에 필요한 초기 속성과 의존성을 설정합니다.
        
        Args:
        options (dict[str, Any] | None):
        역할: `__init__` 실행에 필요한 입력값입니다.
        값: 타입 힌트 기준 `dict[str, Any] | None` 값이 전달됩니다.
        전달 출처: `INGEST 기본 파서` 상위 호출부에서 전달됩니다.
        주의사항: 형식/범위가 함수 가정과 다르면 런타임 예외가 발생할 수 있습니다.
        
        Returns:
        None: 반환값 없이 내부 상태 업데이트, 로깅, 파일 저장 같은 부수효과를 수행합니다.
        
        Raises:
        Exception: 명시적 raise가 없어도 파일 I/O, 네트워크, 외부 라이브러리 예외가 전파될 수 있습니다.
        """

        self.options = options or {}

        self.text_encodings = self._resolve_text_encodings(self.options.get("text_encodings"))

    def parse(self, document: DocumentItem) -> ParsedDocument:
        """
        역할:
        INGEST 기본 파서 문맥에서 `parse` 기능을 수행합니다.
        
        Args:
        document (DocumentItem):
        역할: `parse` 실행에 필요한 입력값입니다.
        값: 타입 힌트 기준 `DocumentItem` 값이 전달됩니다.
        전달 출처: `INGEST 기본 파서` 상위 호출부에서 전달됩니다.
        주의사항: 형식/범위가 함수 가정과 다르면 런타임 예외가 발생할 수 있습니다.
        
        Returns:
        ParsedDocument: INGEST 기본 파서 계산 결과를 `ParsedDocument` 타입으로 반환합니다.
        
        Raises:
        Exception: 명시적 raise가 없어도 파일 I/O, 네트워크, 외부 라이브러리 예외가 전파될 수 있습니다.
        """

        source_path = Path(document.source_path)

        raw_text = self.extract_text(document=document, source_path=source_path)

        metadata = self._build_metadata(document=document, raw_text=raw_text)
        return ParsedDocument(

            document_id=document.document_id,

            source_path=document.source_path,

            file_type=document.file_type,

            raw_text=raw_text,

            metadata=metadata,

        )

    @abstractmethod

    def extract_text(self, document: DocumentItem, source_path: Path) -> str:
        """
        역할:
        INGEST 기본 파서에서 필요한 핵심 텍스트/필드를 추출합니다.
        
        Args:
        document (DocumentItem):
        역할: `extract_text` 실행에 필요한 입력값입니다.
        값: 타입 힌트 기준 `DocumentItem` 값이 전달됩니다.
        전달 출처: `INGEST 기본 파서` 상위 호출부에서 전달됩니다.
        주의사항: 형식/범위가 함수 가정과 다르면 런타임 예외가 발생할 수 있습니다.
        source_path (Path):
        역할: 파일 또는 디렉터리 경로를 지정합니다.
        값: 타입 힌트 기준 `Path` 값이 전달됩니다.
        전달 출처: config 해석 결과 또는 이전 단계 계산 결과가 전달됩니다.
        주의사항: 상대 경로 기준이 다르면 의도하지 않은 위치를 참조할 수 있습니다.
        
        Returns:
        str: 텍스트/프롬프트/SQL 문자열 결과를 반환합니다.
        
        Raises:
        Exception: 명시적 raise가 없어도 파일 I/O, 네트워크, 외부 라이브러리 예외가 전파될 수 있습니다.
        """

    def _parse_text_file(self, source_path: Path) -> str:
        """
        역할:
        INGEST 기본 파서 문맥에서 `_parse_text_file` 기능을 수행합니다.
        
        Args:
        source_path (Path):
        역할: 파일 또는 디렉터리 경로를 지정합니다.
        값: 타입 힌트 기준 `Path` 값이 전달됩니다.
        전달 출처: config 해석 결과 또는 이전 단계 계산 결과가 전달됩니다.
        주의사항: 상대 경로 기준이 다르면 의도하지 않은 위치를 참조할 수 있습니다.
        
        Returns:
        str: 텍스트/프롬프트/SQL 문자열 결과를 반환합니다.
        
        Raises:
        RuntimeError: 내부 검증 실패 또는 하위 호출 오류 상황에서 발생할 수 있습니다.
        """

        last_error: Exception | None = None
        for encoding in self.text_encodings:
            try:
                return source_path.read_text(encoding=encoding)

            except Exception as error:  # pragma: no cover - fallback path

                last_error = error

        if last_error is None:  # pragma: no cover - defensive branch

            raise RuntimeError(f"Failed to parse text file: {source_path}")

        raise RuntimeError(f"Failed to parse text file: {source_path}") from last_error

    @staticmethod

    def _parse_docx_with_python_docx(source_path: Path) -> str:
        """
        역할:
        INGEST 기본 파서 문맥에서 `_parse_docx_with_python_docx` 기능을 수행합니다.
        
        Args:
        source_path (Path):
        역할: 파일 또는 디렉터리 경로를 지정합니다.
        값: 타입 힌트 기준 `Path` 값이 전달됩니다.
        전달 출처: config 해석 결과 또는 이전 단계 계산 결과가 전달됩니다.
        주의사항: 상대 경로 기준이 다르면 의도하지 않은 위치를 참조할 수 있습니다.
        
        Returns:
        str: 텍스트/프롬프트/SQL 문자열 결과를 반환합니다.
        
        Raises:
        ImportError: 내부 검증 실패 또는 하위 호출 오류 상황에서 발생할 수 있습니다.
        """
        try:
            from docx import Document as DocxDocument

        except Exception as error:
            raise ImportError("DOCX 파싱을 위해 python-docx 패키지를 설치하세요.") from error

        doc = DocxDocument(str(source_path))

        paragraphs = [paragraph.text for paragraph in doc.paragraphs if paragraph.text]
        return "\n".join(paragraphs)

    @staticmethod

    def _parse_pdf_with_pypdf(source_path: Path) -> str:
        """
        역할:
        INGEST 기본 파서 문맥에서 `_parse_pdf_with_pypdf` 기능을 수행합니다.
        
        Args:
        source_path (Path):
        역할: 파일 또는 디렉터리 경로를 지정합니다.
        값: 타입 힌트 기준 `Path` 값이 전달됩니다.
        전달 출처: config 해석 결과 또는 이전 단계 계산 결과가 전달됩니다.
        주의사항: 상대 경로 기준이 다르면 의도하지 않은 위치를 참조할 수 있습니다.
        
        Returns:
        str: 텍스트/프롬프트/SQL 문자열 결과를 반환합니다.
        
        Raises:
        ImportError: 내부 검증 실패 또는 하위 호출 오류 상황에서 발생할 수 있습니다.
        """
        try:
            from pypdf import PdfReader

        except Exception as error:
            raise ImportError("PDF 파싱을 위해 pypdf 패키지를 설치하세요.") from error

        reader = PdfReader(str(source_path))

        page_texts: list[str] = []
        for page in reader.pages:
            page_texts.append(page.extract_text() or "")

        return "\n".join(page_texts)

    @staticmethod

    def _coerce_to_text(value: Any) -> str:
        """
        역할:
        INGEST 기본 파서 문맥에서 `_coerce_to_text` 기능을 수행합니다.
        
        Args:
        value (Any):
        역할: `_coerce_to_text` 실행에 필요한 입력값입니다.
        값: 타입 힌트 기준 `Any` 값이 전달됩니다.
        전달 출처: `INGEST 기본 파서` 상위 호출부에서 전달됩니다.
        주의사항: 형식/범위가 함수 가정과 다르면 런타임 예외가 발생할 수 있습니다.
        
        Returns:
        str: 텍스트/프롬프트/SQL 문자열 결과를 반환합니다.
        
        Raises:
        Exception: 명시적 raise가 없어도 파일 I/O, 네트워크, 외부 라이브러리 예외가 전파될 수 있습니다.
        """
        if value is None:
            return ""

        if isinstance(value, str):
            return value

        if isinstance(value, dict):
            for key in ("text", "markdown", "content", "raw_text", "full_text"):
                raw_value = value.get(key)
                if isinstance(raw_value, str) and raw_value.strip():
                    return raw_value

            return str(value)

        if isinstance(value, (list, tuple)):
            text_parts = [BaseParser._coerce_to_text(item) for item in value]
            return "\n".join(part for part in text_parts if part)

        for attr_name in ("text", "markdown", "content", "raw_text", "full_text"):
            attr_value = getattr(value, attr_name, None)
            if isinstance(attr_value, str) and attr_value.strip():
                return attr_value

        return str(value)

    @staticmethod

    def _normalize_metadata(metadata: Any) -> dict[str, Any]:
        """
        역할:
        INGEST 기본 파서 입력값을 표준 형태로 정규화합니다.
        
        Args:
        metadata (Any):
        역할: `_normalize_metadata` 실행에 필요한 입력값입니다.
        값: 타입 힌트 기준 `Any` 값이 전달됩니다.
        전달 출처: `INGEST 기본 파서` 상위 호출부에서 전달됩니다.
        주의사항: 형식/범위가 함수 가정과 다르면 런타임 예외가 발생할 수 있습니다.
        
        Returns:
        dict[str, Any]: INGEST 기본 파서 단계에서 후속 처리에 바로 사용할 매핑 데이터를 반환합니다.
        
        Raises:
        Exception: 명시적 raise가 없어도 파일 I/O, 네트워크, 외부 라이브러리 예외가 전파될 수 있습니다.
        """
        if metadata is None:
            return {}

        if isinstance(metadata, dict):
            return metadata

        return {"_raw_metadata": str(metadata)}

    def _build_metadata(self, document: DocumentItem, raw_text: str) -> dict[str, Any]:
        """
        역할:
        INGEST 기본 파서 문맥에서 `_build_metadata` 기능을 수행합니다.
        
        Args:
        document (DocumentItem):
        역할: `_build_metadata` 실행에 필요한 입력값입니다.
        값: 타입 힌트 기준 `DocumentItem` 값이 전달됩니다.
        전달 출처: `INGEST 기본 파서` 상위 호출부에서 전달됩니다.
        주의사항: 형식/범위가 함수 가정과 다르면 런타임 예외가 발생할 수 있습니다.
        raw_text (str):
        역할: `_build_metadata` 실행에 필요한 입력값입니다.
        값: 타입 힌트 기준 `str` 값이 전달됩니다.
        전달 출처: `INGEST 기본 파서` 상위 호출부에서 전달됩니다.
        주의사항: 형식/범위가 함수 가정과 다르면 런타임 예외가 발생할 수 있습니다.
        
        Returns:
        dict[str, Any]: INGEST 기본 파서 단계에서 후속 처리에 바로 사용할 매핑 데이터를 반환합니다.
        
        Raises:
        Exception: 명시적 raise가 없어도 파일 I/O, 네트워크, 외부 라이브러리 예외가 전파될 수 있습니다.
        """

        base_metadata = self._normalize_metadata(document.metadata)
        return {

            **base_metadata,

            "file_name": document.file_name,

            "file_size": document.file_size,

            "char_length": len(raw_text),

            "selected_parser": self.parser_name,

            "parser_options": self.options,

        }

    @staticmethod

    def _resolve_text_encodings(text_encodings: Any) -> list[str]:
        """
        역할:
        INGEST 기본 파서에서 설정값을 검증 가능한 최종 값으로 확정합니다.
        
        Args:
        text_encodings (Any):
        역할: `_resolve_text_encodings` 실행에 필요한 입력값입니다.
        값: 타입 힌트 기준 `Any` 값이 전달됩니다.
        전달 출처: `INGEST 기본 파서` 상위 호출부에서 전달됩니다.
        주의사항: 형식/범위가 함수 가정과 다르면 런타임 예외가 발생할 수 있습니다.
        
        Returns:
        list[str]: INGEST 기본 파서 결과를 순회 가능한 목록으로 반환합니다.
        
        Raises:
        Exception: 명시적 raise가 없어도 파일 I/O, 네트워크, 외부 라이브러리 예외가 전파될 수 있습니다.
        """
        if isinstance(text_encodings, list):
            normalized = [str(encoding) for encoding in text_encodings if str(encoding).strip()]
            if normalized:
                return normalized

        return ["utf-8", "utf-8-sig", "cp949", "euc-kr"]
