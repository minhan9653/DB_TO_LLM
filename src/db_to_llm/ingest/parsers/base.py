# 이 파일은 모든 파서 구현체가 따라야 하는 기본 인터페이스를 정의한다.
# extract_text()를 추상 메서드로 강제해 파서 교체 시 다른 코드를 건드리지 않아도 된다.
# 텍스트 파일(.sql, .txt, .md)은 BaseParser에 기본 구현이 있어 별도 파서가 필요 없다.
# 하위 파서는 extract_text()만 구현하면 나머지 메타데이터 처리를 공통으로 상속받는다.

from __future__ import annotations

from abc import ABC, abstractmethod
from pathlib import Path
from typing import Any

from src.db_to_llm.ingest.models import DocumentItem, ParsedDocument

# 별도 파서 없이 직접 읽을 수 있는 텍스트 파일 확장자 목록
TEXT_FILE_EXTENSIONS: set[str] = {".sql", ".txt", ".md"}


class BaseParser(ABC):
    """모든 파서 구현체가 상속해야 하는 기본 클래스."""

    parser_name: str = "base"

    def __init__(self, options: dict[str, Any] | None = None) -> None:
        """
        파서 옵션을 초기화한다.

        Args:
            options: 파서별 설정 dict. None이면 빈 dict로 초기화한다.
        """
        self.options = options or {}
        self.text_encodings = self._resolve_text_encodings(self.options.get("text_encodings"))

    def parse(self, document: DocumentItem) -> ParsedDocument:
        """
        문서를 파싱해 ParsedDocument를 반환한다.
        extract_text()를 호출하고 메타데이터를 조립한다.

        Args:
            document: 파싱할 문서의 메타데이터를 담은 DocumentItem.

        Returns:
            ParsedDocument: 추출된 텍스트와 메타데이터를 담은 파싱 결과.
        """
        source_path = Path(document.source_path)
        raw_text = self.extract_text(document=document, source_path=source_path)
        metadata = self._build_metadata(document=document, raw_text=raw_text)
        return ParsedDocument(
            document_id=document.document_id,
            source_path=document.source_path,
            file_type=document.file_type,
            raw_text=raw_text,
            metadata=metadata,
        )

    @abstractmethod
    def extract_text(self, document: DocumentItem, source_path: Path) -> str:
        """
        파일에서 텍스트를 추출한다. 모든 구체 파서가 이 메서드를 구현해야 한다.

        Args:
            document: 문서 메타데이터.
            source_path: 파일 경로 객체.

        Returns:
            str: 추출된 텍스트.
        """
        ...

    def _build_metadata(self, document: DocumentItem, raw_text: str) -> dict[str, Any]:
        """파싱 결과에 포함할 기본 메타데이터를 구성한다."""
        return {
            "parser": self.parser_name,
            "char_count": len(raw_text),
            "file_name": document.file_name,
            "file_type": document.file_type,
            **document.metadata,
        }

    def read_text_file(self, source_path: Path) -> str:
        """텍스트 파일을 여러 인코딩으로 시도해 읽는다."""
        return self._parse_text_file(source_path)

    def _parse_text_file(self, source_path: Path) -> str:
        """설정된 인코딩 순서로 텍스트 파일을 읽는다."""
        last_error: Exception | None = None
        for encoding in self.text_encodings:
            try:
                return source_path.read_text(encoding=encoding)
            except Exception as error:
                last_error = error
        if last_error is None:
            raise RuntimeError(f"텍스트 파일 파싱 실패: {source_path}")
        raise RuntimeError(f"텍스트 파일 파싱 실패: {source_path}") from last_error

    @staticmethod
    def _parse_docx_with_python_docx(source_path: Path) -> str:
        """python-docx를 사용해 DOCX 파일에서 텍스트를 추출한다."""
        try:
            from docx import Document as DocxDocument
        except Exception as error:
            raise ImportError("DOCX 파싱을 위해 python-docx 패키지를 설치하세요.") from error
        doc = DocxDocument(str(source_path))
        paragraphs = [p.text for p in doc.paragraphs if p.text]
        return "\n".join(paragraphs)

    @staticmethod
    def _parse_pdf_with_pypdf(source_path: Path) -> str:
        """pypdf를 사용해 PDF 파일에서 텍스트를 추출한다."""
        try:
            from pypdf import PdfReader
        except Exception as error:
            raise ImportError("PDF 파싱을 위해 pypdf 패키지를 설치하세요.") from error
        reader = PdfReader(str(source_path))
        return "\n".join(page.extract_text() or "" for page in reader.pages)

    @staticmethod
    def _coerce_to_text(value: Any) -> str:
        """임의 값을 텍스트 문자열로 변환한다."""
        if value is None:
            return ""
        if isinstance(value, str):
            return value
        if isinstance(value, dict):
            for key in ("text", "markdown", "content", "raw_text", "full_text"):
                raw_value = value.get(key)
                if isinstance(raw_value, str) and raw_value.strip():
                    return raw_value
            return str(value)
        if isinstance(value, (list, tuple)):
            parts = [BaseParser._coerce_to_text(item) for item in value]
            return "\n".join(p for p in parts if p)
        for attr_name in ("text", "markdown", "content", "raw_text", "full_text"):
            attr_value = getattr(value, attr_name, None)
            if isinstance(attr_value, str) and attr_value.strip():
                return attr_value
        return str(value)

    @staticmethod
    def _normalize_metadata(metadata: Any) -> dict[str, Any]:
        """메타데이터를 dict로 정규화한다."""
        if metadata is None:
            return {}
        if isinstance(metadata, dict):
            return metadata
        return {"_raw_metadata": str(metadata)}

    @staticmethod
    def _resolve_text_encodings(text_encodings: Any) -> list[str]:
        """options에서 인코딩 목록을 확정한다. 미지정 시 기본값 반환."""
        if isinstance(text_encodings, list):
            normalized = [str(e) for e in text_encodings if str(e).strip()]
            if normalized:
                return normalized
        return ["utf-8", "utf-8-sig", "cp949", "euc-kr"]
